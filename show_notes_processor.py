#!/usr/bin/env python3

"""
Show Notes Processor
--------------------
Sends a podcast transcript to the Claude API and generates:
  1. A formatted .docx with show notes
  2. A social-snippets .md file

Adapted from the "podcast-transcribe-and-summarize" Claude skill for the
Private Equity Funcast.
"""

import os
import re
from typing import Any

import anthropic
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

DEFAULT_MODEL = "claude-sonnet-4-20250514"
MAX_TOKENS = 8192

# ---------------------------------------------------------------------------
# System prompt (adapted from SKILL.md)
# ---------------------------------------------------------------------------

SYSTEM_PROMPT = r"""You are a podcast transcript processor for the **Private Equity Funcast**, a podcast covering PE industry trends, deal analysis, founder stories, and operational playbooks. The tone is smart-casual — think sharp colleagues having fun, not a boardroom presentation. PE jargon (EBITDA, multiple expansion, platform vs. bolt-on, dry powder, etc.) is normal and should be preserved, not dumbed down.

## Your Task

Given a plain-text transcript, produce two deliverables in a single response using the exact delimited format described below.

## How to Analyze the Transcript

As you read, identify:
- **Speaker labels** — Jim Milbery is always the host. Map any other speakers as guests.
- **Topic shifts** — where the conversation moves to a new subject.
- **Quotable moments** — punchy, insightful, funny, or contrarian lines.
- **Key takeaways** — the 3-5 things that actually matter, not a summary of everything said.
- **Links and references** — companies, books, articles, people, tools mentioned by name.

If timestamps are present, use them. If not, organize topics chronologically and note "(timestamps unavailable)".

## Output Format

You MUST structure your response using these exact delimiters:

===METADATA===
title: [Episode Title — infer from content]
guest: [Name, Title at Company — or "None" for solo episodes]
date: [Date if mentioned, otherwise "Unknown"]
slug: [kebab-case topic slug, e.g. "ai-in-private-equity"]
===END_METADATA===

===SHOW_NOTES_CONTENT===

## Episode Summary

[2-3 paragraphs for the podcast feed.
First paragraph: hook — what this episode is about and why it matters.
Second paragraph: what ground gets covered, key insights.
Third paragraph (optional): who should listen and what they'll walk away with.]

## Show Notes

[Timestamped or topic-ordered sections]
[HH:MM] Topic Name (or just "Topic Name" if no timestamps)
Brief description of what was discussed (2-3 sentences max per topic)

[Repeat for each topic segment]

## Key Takeaways

- [Opinionated takeaway 1 — not "they discussed X" but the actual insight]
- [Takeaway 2]
- [Takeaway 3]
- [Takeaway 4 (optional)]
- [Takeaway 5 (optional)]

## Recommended Written Format

Primary: [Format name] — [Why this works for this episode]
Working title: "[Draft headline]"
Secondary: [Optional alternative format]
Estimated length: [Short / Medium / Long]

Format options: Long-form narrative, Framework/how-to, Q&A/interview, Point/counterpoint, Listicle/roundup, Case study/deep dive, Opinion/hot take

## Notable Quotes

> "[Quote]" — [Speaker]

> "[Quote]" — [Speaker]

[2-3 of the best quotes — the kind someone would screenshot and share]

## Links & References

- [Item 1]
- [Item 2]
[Anything name-dropped: companies, books, tools, people]

## About the Guest

[2-3 sentence bio based on what was discussed. Skip this section entirely if solo episode.]

===END_SHOW_NOTES_CONTENT===

===SOCIAL_SNIPPETS===

# Social Snippets — [Episode Title]

## Pull Quotes
Short, punchy excerpts that work as standalone social posts or audiogram captions.

> "[Quote text]" — Speaker Name
**Context:** One sentence explaining why this matters.

[Repeat for 3-5 best quotes]

## LinkedIn Post (Long-form)
[A ~150-word post written in Jim's voice. Opens with a hook, shares 1-2 key insights from the episode, ends with a question or CTA to listen. Tag the guest if applicable. Direct, smart, occasionally funny. Not corporate. Not "I'm so humbled to share..."]

## Twitter/X Thread

1. [Hook tweet — the most interesting thing from the episode]
2. [Key insight #1]
3. [Key insight #2]
4. [Quote or contrarian take]
5. [CTA — link to episode]

## Audiogram Candidates

1. "[Quote]" — [Speaker] (est. ~XX seconds)
   **Why this works:** [One line on why it's audiogram-worthy]

2. "[Quote]" — [Speaker] (est. ~XX seconds)
   **Why this works:** [One line]

[2-3 quotes that are ~15-30 seconds spoken, self-contained, compelling out of context]

===END_SOCIAL_SNIPPETS===

## Writing Style Rules

- **Episode summary**: Written for a listener deciding whether to hit play. Lead with the hook, not the guest's bio. Conversational, a little irreverent, but substantive.
- **Show notes**: Scannable. Short descriptions, not transcription.
- **Key takeaways**: Opinionated. Not "they discussed valuation" but "Most PE firms are overpaying for AI companies — and the ones that aren't are using a surprisingly simple framework."
- **Notable quotes**: Punchy > comprehensive. Screenshot-worthy.
- **LinkedIn**: Jim's primary platform. Direct, smart, occasionally funny. Not corporate.
- **Twitter/X**: Punchier. Each tweet should provoke or inform. No filler.
- **Pull quotes**: Must pass the "would someone screenshot this?" test.
- **Audiogram candidates**: Self-contained — must make sense and be interesting with no other context.

## Edge Cases

- **No clear guest**: Drop "About the Guest" and guest-specific social content.
- **Multiple guests**: List all guests, attribute quotes correctly.
- **Very long transcripts**: Focus on the most substantive segments.
- **Poor transcript quality**: Flag garbled sections rather than guessing.
"""


# ---------------------------------------------------------------------------
# Claude API interaction
# ---------------------------------------------------------------------------

def generate_show_notes(transcript_text: str, api_key: str,
                        model: str = DEFAULT_MODEL) -> dict[str, Any]:
    """Send transcript to Claude and return parsed sections.

    Returns a dict with keys:
        metadata  – dict with title, guest, date, slug
        show_notes – raw markdown string of show notes content
        social_snippets – raw markdown string of social snippets
    """
    client = anthropic.Anthropic(api_key=api_key)

    message = client.messages.create(
        model=model,
        max_tokens=MAX_TOKENS,
        system=SYSTEM_PROMPT,
        messages=[
            {
                "role": "user",
                "content": (
                    "Here is the podcast transcript to process:\n\n"
                    + transcript_text
                ),
            }
        ],
    )

    response_text = message.content[0].text
    return parse_response(response_text)


def parse_response(response_text: str) -> dict[str, Any]:
    """Parse Claude's delimited response into structured sections."""

    def extract_section(text: str, start_tag: str, end_tag: str) -> str:
        pattern = re.escape(start_tag) + r"\s*\n(.*?)\n\s*" + re.escape(end_tag)
        match = re.search(pattern, text, re.DOTALL)
        if match:
            return match.group(1).strip()
        return ""

    # Parse metadata
    meta_raw = extract_section(response_text, "===METADATA===", "===END_METADATA===")
    metadata: dict[str, str] = {}
    for line in meta_raw.split("\n"):
        line = line.strip()
        if ":" in line:
            key, _, value = line.partition(":")
            metadata[key.strip().lower()] = value.strip()

    # Parse show notes content
    show_notes = extract_section(
        response_text, "===SHOW_NOTES_CONTENT===", "===END_SHOW_NOTES_CONTENT==="
    )

    # Parse social snippets
    social_snippets = extract_section(
        response_text, "===SOCIAL_SNIPPETS===", "===END_SOCIAL_SNIPPETS==="
    )

    return {
        "metadata": metadata,
        "show_notes": show_notes,
        "social_snippets": social_snippets,
        "raw_response": response_text,
    }


# ---------------------------------------------------------------------------
# .docx generation
# ---------------------------------------------------------------------------

def _parse_show_notes_sections(show_notes_md: str) -> dict[str, str]:
    """Split show notes markdown into sections by ## headings."""
    sections: dict[str, str] = {}
    current_key = ""
    current_lines: list[str] = []

    for line in show_notes_md.split("\n"):
        if line.startswith("## "):
            # Save previous section
            if current_key:
                sections[current_key] = "\n".join(current_lines).strip()
            current_key = line[3:].strip()
            current_lines = []
        else:
            current_lines.append(line)

    # Save last section
    if current_key:
        sections[current_key] = "\n".join(current_lines).strip()

    return sections


def save_show_notes_docx(parsed: dict[str, Any], output_path: str) -> str:
    """Build and save a formatted .docx from parsed show notes.

    Returns the path to the saved file.
    """
    metadata = parsed["metadata"]
    sections = _parse_show_notes_sections(parsed["show_notes"])

    doc = Document()

    # Page setup — US Letter, 1-inch margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # --- Episode Title ---
    title = metadata.get("title", "Untitled Episode")
    title_para = doc.add_heading(title, level=1)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # --- Guest info ---
    guest = metadata.get("guest", "")
    if guest and guest.lower() != "none":
        guest_para = doc.add_paragraph()
        run = guest_para.add_run(f"Guest: {guest}")
        run.italic = True
        run.font.size = Pt(11)

    # --- Date ---
    date = metadata.get("date", "")
    if date and date.lower() != "unknown":
        date_para = doc.add_paragraph()
        run = date_para.add_run(f"Recorded: {date}")
        run.italic = True
        run.font.size = Pt(11)

    doc.add_paragraph("")  # spacer

    # --- Sections ---
    section_order = [
        "Episode Summary",
        "Show Notes",
        "Key Takeaways",
        "Recommended Written Format",
        "Notable Quotes",
        "Links & References",
        "About the Guest",
    ]

    for section_name in section_order:
        content = sections.get(section_name, "")
        if not content:
            continue

        doc.add_heading(section_name, level=2)
        _add_section_content(doc, section_name, content)

    doc.save(output_path)
    return output_path


def _add_section_content(doc: Document, section_name: str, content: str) -> None:
    """Add formatted content for a section to the document."""

    if section_name in ("Key Takeaways", "Links & References"):
        # Bullet-point sections
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue
            # Strip leading bullet markers
            line = re.sub(r"^[-•*]\s*", "", line)
            if line:
                doc.add_paragraph(line, style="List Bullet")

    elif section_name == "Notable Quotes":
        # Quote blocks
        lines = content.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if line.startswith(">"):
                # Extract quote text
                quote_text = line.lstrip("> ").strip()
                para = doc.add_paragraph()
                run = para.add_run(quote_text)
                run.italic = True
                run.font.size = Pt(11)
            elif line and not line.startswith(">"):
                # Attribution or other text
                doc.add_paragraph(line)
            i += 1

    elif section_name == "Show Notes":
        # Timestamped topics — look for [HH:MM] or topic headers
        lines = content.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                i += 1
                continue

            # Check for timestamp pattern or topic header
            ts_match = re.match(r"^\[?(\d{1,2}:\d{2})\]?\s*(.*)", line)
            if ts_match:
                # Topic with timestamp
                timestamp, topic = ts_match.groups()
                para = doc.add_paragraph()
                ts_run = para.add_run(f"[{timestamp}] ")
                ts_run.bold = True
                ts_run.font.size = Pt(11)
                topic_run = para.add_run(topic)
                topic_run.bold = True
                topic_run.font.size = Pt(11)
            elif line and not line.startswith("-") and not line.startswith("•"):
                # Could be a topic name without timestamp
                # Check if it's a short-ish line (likely a header)
                if len(line) < 100 and not line.endswith("."):
                    para = doc.add_paragraph()
                    run = para.add_run(line)
                    run.bold = True
                    run.font.size = Pt(11)
                else:
                    doc.add_paragraph(line)
            else:
                # Description line
                clean = re.sub(r"^[-•*]\s*", "", line)
                if clean:
                    doc.add_paragraph(clean)
            i += 1

    elif section_name == "Recommended Written Format":
        # Short structured section — keep as-is with line breaks
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue
            # Bold the label part (everything before the first colon)
            colon_idx = line.find(":")
            if colon_idx > 0 and colon_idx < 30:
                para = doc.add_paragraph()
                label_run = para.add_run(line[:colon_idx + 1])
                label_run.bold = True
                label_run.font.size = Pt(11)
                value_run = para.add_run(line[colon_idx + 1:])
                value_run.font.size = Pt(11)
            else:
                doc.add_paragraph(line)

    else:
        # Default: plain paragraphs (Episode Summary, About the Guest, etc.)
        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue
            para = doc.add_paragraph(line)
            for run in para.runs:
                run.font.size = Pt(11)


# ---------------------------------------------------------------------------
# Social snippets file
# ---------------------------------------------------------------------------

def save_social_snippets_md(parsed: dict[str, Any], output_path: str) -> str:
    """Save social snippets markdown to file.

    Returns the path to the saved file.
    """
    content = parsed["social_snippets"]
    if not content:
        content = "# Social Snippets\n\nNo social snippets were generated.\n"

    with open(output_path, "w") as f:
        f.write(content + "\n")

    return output_path


# ---------------------------------------------------------------------------
# High-level orchestrator
# ---------------------------------------------------------------------------

def process_transcript(transcript_text: str, api_key: str,
                       output_dir: str,
                       model: str = DEFAULT_MODEL) -> dict[str, str]:
    """Full pipeline: send to Claude, generate .docx and .md.

    Returns a dict with:
        docx_path – path to saved .docx
        md_path   – path to saved .md
        title     – episode title
        guest     – guest name
        slug      – topic slug
    """
    # Call Claude API
    parsed = generate_show_notes(transcript_text, api_key, model)

    metadata = parsed["metadata"]
    slug = metadata.get("slug", "episode")
    title = metadata.get("title", "Untitled Episode")
    guest = metadata.get("guest", "")

    # Build output file paths
    docx_filename = f"show-notes-{slug}.docx"
    md_filename = f"social-snippets-{slug}.md"

    docx_path = os.path.join(output_dir, docx_filename)
    md_path = os.path.join(output_dir, md_filename)

    # Generate files
    save_show_notes_docx(parsed, docx_path)
    save_social_snippets_md(parsed, md_path)

    return {
        "docx_path": docx_path,
        "md_path": md_path,
        "title": title,
        "guest": guest,
        "slug": slug,
    }
